from docx import Document
from .base import BaseLoader
from utils.text_clean import clean_text

class DocxLoader(BaseLoader):
    def load(self, path):
        doc = Document(path)
        parts = []

        # Extract all paragraphs from the document
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())

        # Extract table contents
        for table in doc.tables:
            for row in table.rows:
                # Join all cell values in a row with a separator
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        
        # Combine everything into a single cleaned-up string
        return clean_text("\n".join(parts))